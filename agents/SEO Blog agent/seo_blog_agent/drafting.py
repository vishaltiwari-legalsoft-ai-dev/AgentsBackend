"""Team steps 11-12: guidelines + outline + hard rules → the draft, plus the
compliance checklist that proves the spec was followed. No LLM → no draft
(a fabricated draft would be worse than an honest failure)."""
from __future__ import annotations

import io
import re
from pathlib import Path

from seo_geo_agent import sources

GUIDELINES_PATH = Path(__file__).with_name("guidelines.md")


def _hard_rules(sheet: dict, outline_doc: dict) -> str:
    lsi = [l["term"] for l in sheet["lsi"]]
    links_line = ""
    if sheet.get("internal_links"):
        links_line = f"\n- Link these pages of ours where relevant: {sheet['internal_links']}."
    return (
        f'- Use the exact phrase "{sheet["keyword"]}" between {sheet["usage"]["target_min"]} '
        f"and {sheet['usage']['target_max']} times.\n"
        f"- Work every one of these LSI terms in naturally: {lsi}.\n"
        f"- Target length: {outline_doc['targets']['word_count']} words (stay within 10%).\n"
        "- Cite every listed source inside its section, by its specific name, as a markdown link.\n"
        f"- H1: {outline_doc['meta']['title']}. Meta description: {outline_doc['meta']['description']}. "
        f"Slug: {outline_doc['meta']['slug']}." + links_line
    )


def build_draft(sheet: dict, outline_doc: dict, citations: dict, llm=None) -> dict:
    llm = llm or sources.llm_text
    guidelines = GUIDELINES_PATH.read_text(encoding="utf-8")
    outline_lines = "\n".join(
        f"{'#' * o['level']} {o['heading']}  — {o['note']}"
        + (f" [keywords: {', '.join(o['keywords'])}]" if o.get("keywords") else "")
        for o in outline_doc["outline"]
    )
    cited = "\n".join(
        f"- [{c['id']}] {c['source_name']} — {c['claim']} ({c['url']}) → section: {c['section']}"
        for c in citations["items"]
    ) or "- (no verified sources available — do not invent any)"
    markdown = llm(
        "You are Legal Soft's senior SEO content writer. Write the complete article in clean "
        "markdown. Output the article only — no preamble.",
        f"{guidelines}\n\n## Hard rules\n{_hard_rules(sheet, outline_doc)}\n\n"
        f"## Outline (follow exactly)\n{outline_lines}\n\n## Verified sources\n{cited}",
    )
    return {"markdown": markdown, "meta": outline_doc["meta"],
            "compliance": check_compliance(markdown, sheet, outline_doc, citations), "edited": False}


def check_compliance(markdown: str, sheet: dict, outline_doc: dict, citations: dict) -> dict:
    text = markdown.lower()
    words = len(re.findall(r"\w+", markdown))
    lo, hi = sheet["usage"]["target_min"], sheet["usage"]["target_max"]
    count = len(re.findall(re.escape(sheet["keyword"].lower()), text))
    target_wc = outline_doc["targets"]["word_count"]
    lsi_missing = [l["term"] for l in sheet["lsi"] if l["term"].lower() not in text]
    cite_missing = [c["id"] for c in citations["items"] if c["url"] not in markdown]
    heads_missing = [o["heading"] for o in outline_doc["outline"] if o["heading"].lower() not in text]
    meta = outline_doc["meta"]
    checks = [
        {"id": "kw_count", "label": f"Main keyword used {lo}-{hi}×", "pass": lo <= count <= hi,
         "detail": f"found {count}×"},
        {"id": "lsi", "label": f"All {len(sheet['lsi'])} LSI terms present", "pass": not lsi_missing,
         "detail": ("missing: " + ", ".join(lsi_missing[:5])) if lsi_missing else "all present"},
        {"id": "word_count", "label": f"~{target_wc} words (-10%/+25%)",
         "pass": 0.9 * target_wc <= words <= 1.25 * target_wc, "detail": f"{words} words"},
        {"id": "citations", "label": "Every verified citation linked", "pass": not cite_missing,
         "detail": ("missing: " + ", ".join(cite_missing)) if cite_missing else
                   f"{len(citations['items'])} linked"},
        {"id": "sections", "label": "Every outline section present", "pass": not heads_missing,
         "detail": ("missing: " + "; ".join(heads_missing[:4])) if heads_missing else "all present"},
        {"id": "meta", "label": "Meta title + description attached",
         "pass": bool(meta["title"] and meta["description"]), "detail": meta["title"]},
    ]
    return {"checks": checks, "all_pass": all(c["pass"] for c in checks)}


_MD_LINK = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")


def to_docx(markdown: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        flat = _MD_LINK.sub(r"\1 (\2)", stripped)
        m = re.match(r"^(#{1,4})\s+(.*)", flat)
        if m:
            doc.add_heading(m.group(2), level=len(m.group(1)))
        elif flat.startswith(("- ", "* ")):
            doc.add_paragraph(flat[2:], style="List Bullet")
        else:
            doc.add_paragraph(flat)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
